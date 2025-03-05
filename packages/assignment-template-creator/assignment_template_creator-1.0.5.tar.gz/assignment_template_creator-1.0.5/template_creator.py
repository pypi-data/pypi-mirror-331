import os
from docx import Document
from pathlib import Path
from string import Template
from colorama import Fore, Back, Style
from datetime import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# report template
MD_REPORT_TEMPLATE = Template("""<div align="center">
# ${module_code} W${week_number} Assignment ${report_type}
## ${title}
### Tutor: ${tutor_name}
### Student: ${matriculation}
### Date: ${date}
</div>

---

## **Introduction**
Provide a brief overview of the topic, objectives, and scope of the assignment.

---

## **Design**
Explain your approach, methodology, and design decisions. Attempt to justify your questionable structural decisions.

---

## **Testing**
Explain how you tested your program.

### **Test Results**
| **Test Name** | **Expected Output**  |  **Actual Output**   | **Result** |
|-----------------|----------------------|----------------------|------------|
|                 |                      |                      |            |
|                 |                      |                      |            |
|                 |                      |                      |            |


---

## **Evaluation**
Evaluate the success of your program against what you were asked to do.

---

## **Conclusion**
Summarise key points, pretend you genuinely enjoyed the process of creating this, and talk about potential imporvements you'll never implemnent.

---

### **References** *(if needed)*
Include citations in academic format.
""")

PRINT_TEST = Template("""import os
import subprocess
import filecmp
import re
from docx import Document

TEST_DIR = "Custom-Tests"
language = "${language}"
report_extension = "${report_extension}"

COMPILATIONS = {
    "c": "make",
    "c++": "make",
    "java": "chmod +x runJava.sh",
    "haskell": "chmod +x runHaskell.sh",
    "python": "chmod +x run_python.sh",
    "javascript": "chmod +x runJS.sh"
}

RUNS = {
    "c": "./main",
    "c++": "./main",
    "java": "./runJava.sh",
    "haskell": "./runHaskell.sh",
    "python": "./run_python.sh",
    "javascript": "./runJS.sh"
}


def find_tests():
    test_cases = []
    for file in os.listdir(TEST_DIR):
        if file.endswith(".in"):
            test_name = file[:-3]  # file name without extension
            expected = f"{test_name}.out"
            if expected in os.listdir(TEST_DIR):
                test_cases.append(test_name)
    return test_cases

def compile_code():
    if language in COMPILATIONS:
        COMP = COMPILATIONS.get(language)
        try:
            subprocess.run(COMP, shell=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=False)
        except Exception as error:
            print(f"COMPILATION FAILED: {error}")
            return 1
    else:
        print("This language is not supported. Modify COMPILATIONS and RUNS in run_tests.py.")
        return 1

def run_tests(test):
    input_file_p = os.path.join(TEST_DIR, f"{test}.in")
    output_file_p = os.path.join(TEST_DIR, f"{test}.out")
    actual_output_p = os.path.join(TEST_DIR, f"{test}.act")

    EXECUTABLE = RUNS.get(language)

    try:
        with open(input_file_p, "r") as input_file, open(actual_output_p, "w") as actual_output:
            subprocess.run(EXECUTABLE, stdin=input_file, stdout=actual_output, stderr=actual_output, check=False)

        if filecmp.cmp(output_file_p, actual_output_p):
            print(f"* COMPARISON TEST - {TEST_DIR}/{test}.out : pass")
            return 1
        else:
            print(f"* COMPARISON TEST - {TEST_DIR}/{test}.out : fail")
            print("--- expected output ---")
            with open(output_file_p) as file:
                print(file.read())
            print("--- actual output ---")
            with open(actual_output_p) as file:
                print(file.read())
            return 0
    except Exception as error:
        print(f"ERROR in subprocess: {error}")
        return 1

def main():
    test_cases = find_tests()
    cnt = 0
    compile_code()
    for test in test_cases:
        cnt += run_tests(test)

    all_tests = len(test_cases)
    print(f"{cnt} out of {all_tests} tests passed.")

if __name__ == "__main__":
    main()
""")

LINK_TEST = Template(r"""import os
import subprocess
import filecmp
import re
from docx import Document

TEST_DIR = "Custom-Tests"
language = "${language}"
report_extension = "${report_extension}"
report_file = "${report_path}"

COMPILATIONS = {
    "c": "make",
    "c++": "make",
    "java": "chmod +x runJava.sh",
    "haskell": "chmod +x runHaskell.sh",
    "python": "chmod +x run_python.sh",
    "javascript": "chmod +x runJS.sh"
}

RUNS = {
    "c": "./main",
    "c++": "./main",
    "java": "./runJava.sh",
    "haskell": "./runHaskell.sh",
    "python": "./run_python.sh",
    "javascript": "./runJS.sh"
}


def find_tests():
    test_cases = []
    for file in os.listdir(TEST_DIR):
        if file.endswith(".in"):
            test_name = file[:-3]  # file name without extension
            expected = f"{test_name}.out"
            if expected in os.listdir(TEST_DIR):
                test_cases.append(test_name)
    return test_cases

def compile_code():
    if language in COMPILATIONS:
        COMP = COMPILATIONS.get(language)
        try:
            subprocess.run(COMP, shell=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=False)
        except Exception as error:
            print(f"COMPILATION FAILED: {error}")
            return 1
    else:
        print("This language is not supported. Modify COMPILATIONS and RUNS in run_tests.py.")
        return 1

def run_tests(test):
    input_file_p = os.path.join(TEST_DIR, f"{test}.in")
    output_file_p = os.path.join(TEST_DIR, f"{test}.out")
    actual_output_p = os.path.join(TEST_DIR, f"{test}.act")

    EXECUTABLE = RUNS.get(language)

    try:
        with open(input_file_p, "r") as input_file, open(actual_output_p, "w") as actual_output:
            subprocess.run(EXECUTABLE, shell=True, stdin=input_file, stdout=actual_output, stderr=actual_output, check=False)

        with open(output_file_p, "r") as file1, open(actual_output_p, "r") as file2:
            expected_output = file1.read().strip()
            actual_output = file2.read().strip()

        if expected_output == actual_output:
            return test, expected_output, actual_output, "PASS"
        else:
            return test, expected_output, actual_output, "FAIL"

    except Exception as error:
        print(f"ERROR in subprocess: {error}")
        return test, "ERROR", "ERROR", "FAIL"

def update_md(results):
    if os.path.exists(report_file):
        with open(report_file, "r") as file:
            lines = file.readlines()
    else:
        lines = []

    new_lines = []
    table = False
    for line in lines:
        if line.startswith("| **Test Name** |"):
            table = True
            # add updated table
            new_lines.append("\n| **Test Name** | **Expected Output** | **Actual Output** | **Resut** |\n")
            new_lines.append("|-----------|----------------|--------------|--------|\n")
            for test, expected, actual, result in results:
                new_lines.append(f"| {test} | {expected} | {actual} | {result} |\n")
            continue # skip table header
        elif table and line.strip() == "":
            table = False # end of table
        if not table:
            new_lines.append(line) # keep all the report contents

    # write back to he report
    with open(report_file, "w") as file:
        file.writelines(new_lines)

def update_docx(results):
    try:
        doc = Document(report_file)
    except FileNotFoundError:
        print("File not found. creating a new report")
        doc = Document()

    test_table = None
    for table in doc.tables:
        if "Test Name" in table.rows[0].cells[0].text:
            test_table = table
            break

    if test_table:
        for _ in range(len(test_table.rows) - 1):
            test_table._element.remove(test_table.rows[1]._element)
    else:
        test_table = doc.add_table(rows=1, cols=4)
        test_table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Test Name"
        hdr_cells[1].text = "Expected Outcome"
        hdr_cells[2].text = "Actual Outcome"
        hdr_cells[3].text = "Result"

    for test, expected, actual, result in results:
        row_cells = table.add_row().cells
        row_cells[0].text = test
        row_cells[1].text = expected
        row_cells[2].text = actual
        row_cells[3].text = result

    doc.save(report_file)

def main():
    test_cases = find_tests()
    if not test_cases:
        print ("ERROR : no test cases found.")
        exit(1)

    compile_code()

    results = [run_tests(test) for test in test_cases]

    if report_extension == "md":
        update_md(results)
    else:
        update_docx(results)


if __name__ == "__main__":
    main()
""")

def create_project():
    # get user input
    print(Fore.LIGHTBLUE_EX + " Welcome to Ace-ssignment template -- the ultimate CS assignemnt template !\n Enter your coursework details below" + Style.RESET_ALL)
    project_type = input_loop(
        "project type (Group (g) / Individual (i))", {"g", "i"}, 1)
    week_number = input_loop("week number (optional)", {"any"}, 2)
    title = input_loop("project title (optional)", {"any"}, 3)
    print(Fore.LIGHTBLACK_EX +
          "    Ace-ssignment template provides basic project setups for the following languages :")
    print(Fore.LIGHTBLUE_EX + "         Java, Pyhton, C, C++, Haskell, and JavaScript" +
          Fore.LIGHTBLACK_EX)
    print("    Feel free to skip that part of the setup by pressing ENTER." + Style.RESET_ALL)
    language = input_loop("programming language", {"java", "python", "c",
                                                   "c++", "haskell",
                                                   "javascript", ""}, 4)

    # report details input
    module_code = input_loop("module code (optional)", {"any"}, 5)
    tutor_name = input_loop("tutor name (optional)", {"any"}, 6)
    matriculation = input_loop("matriculation number (oprional)", {"any"}, 7)
    report_format = input_loop("report format (Markdown (m) / Word (w))", {"w", "m"}, 8)
    print(Fore.LIGHTBLACK_EX +
          "Ace-ssignment template provides two automated testing options :")
    print(Fore.LIGHTBLUE_EX + "--> Automatic Testing Table Generation :" + Fore.BLUE + "\n    your testing table will be populated automatically by your test names, expected outputs, and actual outputs, \n    extracted form the Custom-Tests directory, which you can fill with your own .in and .out files. \n    The testing table can be updated by running `python3 run_tests.py`" + Fore.LIGHTBLUE_EX + " \n\n--> Custom stacscheck Output :" + Fore.BLUE + "\n    running `python3 run_tests.py` will print out a stacscheck-style output\n    based on the contents of your Custom-Tests foler." +
          Fore.LIGHTBLACK_EX)
    print("    Feel free to skip that part of the setup by pressing ENTER.\n    Note : if you have not selected a language, you may have to tweak `python3 run_tests.py` to support your language." + Style.RESET_ALL)
    test = input_loop("test setting (automatic testing table (a) / custom stacscheck output (c))", {"a", "c", ""}, 9)

    # folder name and base path based on input and folder name
    folder_name = f"W{week_number}-Assignment"
    base_path = Path(folder_name)

    if report_format == "m":
        report_extension = "md"
    else:
        report_extension = "docx"

    if project_type == "i":  # individual report setup
        src_path = base_path / "src"
        report_path = "" # not used but passed in as argument
    else:  # group project setup
        src_path = base_path / "Code" / "src"
        report_path = base_path / "Reports"

    # programming language specifics
    if language == "java":
        main_file = src_path / "Main.java"
        main_file_content = """public class Main {
    public static void main(String args[]) {
        System.out.print("Good luck with your coursework ;)");
    }
}"""
        script_content = """#!/bin/bash
# this is a script to run your java project
# add files to the compile and run sections as needed

# compile your files
javac Main.java

# run your files
java Main"""
        script_file = src_path / "runJava.sh"

    elif language == "c":
        main_file = src_path / "main.c"
        main_file_content = """#include <stdio.h>

int main(int argc, char **argv) {
    printf("Good luck with your coursework ;)");
    return 0;
}"""
        script_content = """CC = gcc
TARGET = main
CFLAGS = -Wall -Wextra
OBJS = main.o

$(TARGET): $(OBJS)
\t$(CC) $(CFLAGS) -o $(TARGET) $(OBJS)

main.o: main.c
\t$(CC) $(CFLAGS) -c main.c

clean:
\trm -f *.o $(TARGET)"""
        script_file = src_path / "Makefile"

    elif language == "haskell":
        main_file = src_path / "Main.hs"
        main_file_content = """main :: IO ()
main = putStr "Good luck with your coursework ;)"
        """
        script_file = src_path / "runHaskell.sh"
        script_content = """#!/bin/bash

# this is a script to run your Haskell project
# add files to the compile and run sections as needed

# compile your files
ghc -o main Main.hs

# run your files
./main
        """

    elif language == "javascript":
        main_file = src_path / "main.js"
        main_file_content = """console.log("Good luck with your coursework ;)");"""
        script_file = src_path / "runJS.sh"
        script_content = """#!/bin/bash
# this is a script to run your JavaScript project
# add files as needed

# run your files
node main.js
"""

    elif language == "python":
        main_file = src_path / "main.py"
        main_file_content = """if __name__ == "__main__":
    print("Good luck with your coursework ;)")"""
        script_file = src_path / "run_python.sh"
        script_content = """#!/bin/bash
# this is a script to run your Python project
# add files as needed

# run your files
python3 main.py
"""

    elif language == "c++":
        main_file = src_path / "main.cpp"
        main_file_content = """#include <iostream>
int main() {
    std::cout << "Good luck with your coursework ;)";
    return 0;
}
"""
        script_file = src_path / "Makefile"
        script_content = """CXX = g++
TARGET = main
CXXFLAGS = -Wall -Wextra
OBJS = main.o

$(TARGET): $(OBJS)
\t$(CXX) $(CXXFLAGS) -o $(TARGET) $(OBJS)

main.o: main.cpp
\t$(CXX) $(CXXFLAGS) -c main.cpp

clean:
\trm -f *.o $(TARGET)"""

    # create directories using Path.mkdir
    src_path.mkdir(parents=True, exist_ok=True)

    # configure project files
    if language != "":
        main_file.write_text(main_file_content)
        script_file.write_text(script_content)

    date=str({datetime.today().strftime('%d %m %Y')})
    date = date[2:len(date)-2]
    test_report = make_report(report_extension, project_type, week_number, module_code, tutor_name, title, matriculation, date, base_path, report_path)

    if test != "":
        if project_type == "i":
            make_test_script(language, report_extension, src_path, test_report, test)
        else:
            make_test_script(language, report_extension, src_path, test_report, test)

    print(Fore.GREEN + "All done :з project created !" + Style.RESET_ALL)


def input_loop(prompt, target_set, num):
    if "any" in target_set:
        return input(f" {num}. Enter {prompt} : ").strip()

    inp = input(f" {num}. Enter {prompt} : ").strip().lower()
    while inp not in target_set:
        print(Fore.RED + "Invalid input" + Style.RESET_ALL)
        inp = input(f" {num}. Enter {prompt} : ").strip().lower()
    return inp

def make_test_script(language, report_extension, src_path, report_path, test):
    # Supported languages and their execution commands
    test_script = src_path / "run_tests.py"
    test_dir = src_path / "Custom-Tests"
    test_dir.mkdir(parents=True, exist_ok=True)

    test_in1 = test_dir / "test1.in"
    test_out1 = test_dir / "test1.out"
    test_in1.write_text("")
    test_out1.write_text("Good luck with your coursework ;)")

    test_in2 = test_dir / "test2.in"
    test_out2 = test_dir / "test2.out"
    test_in2.write_text("")
    test_out2.write_text("Hope you fail >:(")

    if test == "a":
        test_script_contents = LINK_TEST.substitute(language=language, report_extension=report_extension, report_path=report_path)
    else:
        test_script_contents = PRINT_TEST.substitute(language=language, report_extension=report_extension, report_path=report_path)
    test_script.write_text(test_script_contents)


def make_report(report_extension, project_type, week_number, module_code, tutor_name, title, matriculation, date, base_path, report_path):
    if report_extension == "docx":
        return make_word_report(project_type, week_number, module_code, tutor_name, title, matriculation, date, base_path, report_path)
    else:
        return make_md_report(project_type, week_number, module_code, tutor_name, title, matriculation, date, base_path, report_path)


def make_word_report(project_type, week_number, module_code, tutor_name, title, matriculation, date, base_path, report_path):
    if project_type == "i":
        doc = make_word("", week_number, module_code, tutor_name, title, matriculation, date)
        save_path = Path(base_path) / f"W{week_number}-Report.docx"
        doc.save(save_path)
        return f"../W{week_number}-Report.docx"

    else:
        report_path.mkdir(parents=True, exist_ok=True)
        save_path_ind = Path(report_path) / "Individual-Report.docx"
        save_path_gr = Path(report_path) / "Group-Report.docx"
        doc_ind = make_word("Individual Report", week_number, module_code, tutor_name, title, matriculation, date)
        doc_gr = make_word("Group Report", week_number, module_code, tutor_name, title, matriculation, date)
        doc_ind.save(save_path_ind)
        doc_gr.save(save_path_gr)
        return f"../../Reports/Group-Report.docx"

def make_word(rep_t, week_number, module_code, tutor_name, title, matriculation, date):
        doc = Document()

        heading = doc.add_heading(
            f"{module_code} W{week_number} Assignment {rep_t}", 1)
        head_p = doc.add_paragraph(
            f"Title: {title}\nTutor: {tutor_name}\nStudent: {matriculation}\nDate: {date}\n")

        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        head_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


        doc.add_page_break()

        doc.add_heading("Introduction", 2)
        doc.add_paragraph(
            "Provide a brief overview of the topic, objectives, and scope of the assignment.")

        doc.add_paragraph("Design")
        doc.add_paragraph(
            "Explain your approach, methodology, and design decisions. Attempt to justify your questionable structural decisions.")


        if rep_t == "" or rep_t == "Group Report":
            doc.add_heading("Testing Strategies", 2)
            doc.add_paragraph("Explain how you tested your program.")

            doc.add_paragraph("Test Results")

            table = doc.add_table(rows=3, cols=4)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Test Name"
            hdr_cells[1].text = "Expected Output"
            hdr_cells[2].text = "Actual Output"
            hdr_cells[3].text = "Result"

            for i in range(1, 3):
                row_cells = table.rows[i].cells
                row_cells[0].text = ""
                row_cells[1].text = ""
                row_cells[2].text = ""
                row_cells[3].text = ""

        doc.add_heading("Evaluation", 2)
        doc.add_paragraph("Evaluate the success of your program against what you were asked to do.")

        doc.add_heading("Conclusion", 2)
        doc.add_paragraph("Summarise key points, pretend you genuinely enjoyed the process of creating this, and talk about potential improvements you'll never implement.")

        doc.add_heading("References)", 2)
        doc.add_paragraph("Include citations in academic format.")

        return doc


def make_md_report(project_type, week_number, module_code, tutor_name, title, matriculation, date, base_path, report_path):
    # create folder structure
    if project_type == "i":  # individual report setup
        report_file = base_path / f"W{week_number}-Report.md"

        report_type = ""
        report_contents = MD_REPORT_TEMPLATE.substitute(module_code=module_code, week_number=week_number, report_type=report_type, tutor_name=tutor_name, title=title, matriculation=matriculation, date=date)
        # write report files
        report_file.write_text(report_contents)
        return f"../W{week_number}-Report.md"

    else:  # group project setup
        report_path.mkdir(parents=True, exist_ok=True)
        group_report_file = report_path / "Group-Report.md"
        report_type_g = "Group Report"

        individual_report_file = report_path / "Individual-Report.md"
        report_type_i = "Individual Report"

        report_g_contents = MD_REPORT_TEMPLATE.substitute(module_code=module_code, week_number=week_number, report_type=report_type_g, tutor_name=tutor_name, title=title, matriculation=matriculation, date=date)
        group_report_file.write_text(report_g_contents)

        report_i_contents = MD_REPORT_TEMPLATE.substitute(module_code=module_code, week_number=week_number, report_type=report_type_i, tutor_name=tutor_name, title=title, matriculation=matriculation, date=date)
        individual_report_file.write_text(report_i_contents)
        return f"../../Reports/Group-Report.md"


def main():
    create_project()

if __name__ == "__main__":
    main()
