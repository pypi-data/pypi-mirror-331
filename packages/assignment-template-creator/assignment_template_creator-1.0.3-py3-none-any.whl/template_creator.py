import os
from docx import Document
from pathlib import Path
from string import Template
from colorama import Fore, Back, Style
from datetime import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# report template
MD_REPORT_TEMPLATE = Template("""<div align="center">

# ${module_code} W${week_number} Assignment ${report_type}
## ${title}
### Tutor: ${tutor_name}
### Student: ${matriculation}
### Date: ${date}
</div>

---

## **Introduction**
Provide a brief overview of the topic, objectives, and scope of the assignment.

---

## **Design**
Explain your approach, methodology, and design decisions.

---

## **Testing**

### **Implementation Details**
Describe the development and coding process.

### **Testing Strategies**
Explain how you tested your program.

### **Test Results**
| **Test Description**   | **Expected Result**   | **Actual Result**   |
|------------------------|----------------------|----------------------|
| Test 1:                |                      |                      |
| Test 2:                |                      |                      |
| Test 3:                |                      |                      |


---

## **Conclusion**
Summarise key points, pretend you genuinely enjoyed the process of creating this, and talk about potential imporvements you'll never implemnent.

---

### **References** *(if needed)*
Include citations in academic format.
""")


def create_project():
    # get user input
    print(Fore.LIGHTBLUE_EX + "Welcome to the ultimate CS assignemnt template ! Enter your coursework details below" + Style.RESET_ALL)
    project_type = input_loop(
        "project type (Group (g) / Individual (i))", {"g", "i"})
    week_number = input_loop("week number (optional)", {"any"})
    title = input_loop("project title (optional)", {"any"})
    print(Fore.LIGHTBLACK_EX +
          "\nThis template creator provides basic project setups for the following languages :")
    print(Fore.BLUE + "Java, Pyhton, C, C++, Haskell, and JavaScript" +
          Fore.LIGHTBLACK_EX)
    print("Feel free to skip that part of the setup by pressing ENTER." + Style.RESET_ALL)
    language = input_loop("programming language", {"java", "python", "c",
                                                   "c++", "haskell",
                                                   "javascript", ""})

    # report details input
    module_code = input_loop("module code (optional)", {"any"})
    tutor_name = input_loop("tutor name (optional)", {"any"})
    matriculation = input_loop("matriculation number (oprional)", {"any"})
    report_format = input_loop(
        "report format (Markdown (m) / Word (w))", {"w", "m"})

    # folder name and base path based on input and folder name
    folder_name = f"W{week_number}-Assignment"
    base_path = Path(folder_name)

    if report_format == "m":
        report_extention = "md"
    else:
        report_extention = "docx"

    if project_type == "i":  # individual report setup
        src_path = base_path / "src"
        report_path = "" # not used but passed in as arg
    else:  # group project setup
        src_path = base_path / "Code" / "src"
        report_path = base_path / "Reports"

    # programming language specifics
    if language == "java":
        main_file = src_path / "Main.java"
        main_file_content = """public class Main {
    public static void main(String args[]) {
        System.out.println("Good luck with your coursework ;");
    }
}"""
        script_content = """#!/bin/bash
# this is a script to run your java project
# add files to the compile and run sections as needed

# compile your files
javac main.java

# run your files
java main"""
        script_file = src_path / "runJava.sh"

    elif language == "c":
        main_file = src_path / "Main.c"
        main_file_content = """#include <stdio.h>

int main(int argc, char **argv) {
    printf("Good luck with your coursework ;)\n");
    return 0;
}"""
        script_content = """CC = GCC
TARGET = main
FLAGS = -Wall -Wextra
OBJS = main.o

$(TARGET): $(OBJS)
    $(CC) $(CFLAGS) -o $(TARGET) $(OBJS)

main.o: main.c
    $(CC) $(CFLAGS) -c main.c

clean:
    rm -f *.o $(TARGET)"""
        script_file = src_path / "Makefile"

    elif language == "haskell":
        main_file = src_path / "Main.hs"
        main_file_content = """main :: IO ()
main = putStrLn "Good luck with your coursework ;)"
        """
        script_file = src_path / "runHaskell.sh"
        script_content = """#!/bin/bash

# this is a script to run your Haskell project
# add files to the compile and run sections as needed

# compile your files
ghc -o main Main.hs

# run your files
./main
        """

    elif language == "javascript":
        main_file = src_path / "main.js"
        main_file_content = """console.log("Good luck with your coursework ;)");"""
        script_file = src_path / "runJS.sh"
        script_content = """#!/bin/bash
# this is a script to run your JavaScript project
# add files as needed

# run your files
node main.js
"""

    elif language == "python":
        main_file = src_path / "main.py"
        main_file_content = """if __name__ == "__main__":
    print("Good luck with your coursework ;)")"""
        script_file = src_path / "run_python.sh"
        script_content = """#!/bin/bash
# this is a script to run your Python project
# add files as needed

# run your files
python3 main.py
"""

    elif language == "c++":
        main_file = src_path / "main.cpp"
        main_file_content = """#include <iostream>
int main() {
    std::cout << "Good luck with your coursework ;)" << std::endl;
    return 0;
}
"""
        script_file = src_path / "Makefile"
        script_content = """CXX=g++
TARGET=main
CXXFLAGS = -Wall -Wextra
OBJS = main.o

$(TARGET): $(OBJS)
    $(CXX) $(CXXFLAGS) -o $(TARGET) $(OBJS)

main.o: main.cpp
    $(CXX) $(CXXFLAGS) -c main.cpp

clean:
    rm -f *.o $(TARGET)
"""

    # create directories using Path.mkdir
    src_path.mkdir(parents=True, exist_ok=True)

    make_test_script(language, src_path)

    # configure project files
    if language != "":
        main_file.write_text(main_file_content)
        script_file.write_text(script_content)

    date=str({datetime.today().strftime('%d %m %Y')})
    date = date[2:len(date)-2]
    make_report(report_extention, project_type, week_number, module_code, tutor_name, title, matriculation, date, base_path, report_path)

    print(Fore.GREEN + "All done :з project created !" + Style.RESET_ALL)


def input_loop(prompt, target_set):
    inp = input(f"Enter {prompt} : ").strip().lower()

    if "any" in target_set:
        return inp

    while inp not in target_set:
        print(Fore.RED + "Invalid input" + Style.RESET_ALL)
        inp = input(f"Enter {prompt} : ").strip().lower()
    return inp

def make_test_script(language, src_path):
    # Supported languages and their execution commands
    LANGUAGE_COMMANDS = {
        "c": "./a.out",
        "cpp": "./a.out",
        "java": "java Main",
        "haskell": "./main",
        "python": "python3 main.py",
        "javascript": "node main.js"
    }

    # Compilation commands for compiled languages
    COMPILATION_COMMANDS = {
        "c": "gcc main.c -o a.out",
        "cpp": "g++ main.cpp -o a.out",
        "java": "javac Main.java",
        "haskell": "ghc -o main main.hs"
    }

    test_script = src_path / "run_tests.py"
    test_dir = src_path / "Custom-Tests"
    test_dir.mkdir(parents=True, exist_ok=True)
    test_script_contents = f"""import os
import subprocess
import filecmp
import re
from docx import Document

TEST_DIR = "Custom-Tests"

LANGUAGE_COMMANDS = {{
    "c": "./a.out",
    "cpp": "./a.out",
    "java": "java Main",
    "haskell": "./main",
    "python": "python3 main.py",
    "javascript": "node main.js"
}}

def find_tests():
    test_cases = []
    for file in os.listdir(TEST_DIR):
        if file.endswith(".in"):
            test_name = file[:-3]  # file name without extension
            expected = f"{{test_name}}.out"
            if expected in os.listdir(TEST_DIR):
                test_cases.append(test_name)
    return test_cases

def run_tests(test):
    input_file_p = os.path.join(TEST_DIR, f"{{test}}.in")
    output_file_p = os.path.join(TEST_DIR, f"{{test}}.out")
    actual_output_p = os.path.join(TEST_DIR, f"{{test}}.act")

    EXECUTABLE = LANGUAGE_COMMANDS.get("{language}")

    try:
        with open(input_file_p, "r") as input_file, open(actual_output_p, "w") as actual_output:
            subprocess.run(EXECUTABLE, shell=True, stdin=input_file, stdout=actual_output, stderr=actual_output, check=False)

        if filecmp.cmp(output_file_p, actual_output_p):
            print(f"* COMPARISON TEST - {{TEST_DIR}}/{{test}}.out : pass")
            return 1
        else:
            print(f"* COMPARISON TEST - {{TEST_DIR}}/{{test}}.out : fail")
            print("--- expected output ---")
            with open(output_file_p) as file:
                print(file.read())
            print("--- actual output ---")
            with open(actual_output_p) as file:
                print(file.read())
            return 0
    except Exception as error:
        print(f"ERROR in subprocess: {{error}}")
        return 0

def main():
    test_cases = find_tests()
    cnt = 0

    for test in test_cases:
        cnt += run_tests(test)

    all_tests = len(test_cases)
    print(f"{{cnt}} out of {{all_tests}} tests passed.")

if __name__ == "__main__":
    main()
"""
    test_script.write_text(test_script_contents)


def make_report(report_extention, project_type, week_number, module_code, tutor_name, title, matriculation, date, base_path, report_path):
    if report_extention == "docx":
        make_word_report(project_type, week_number, module_code, tutor_name, title, matriculation, date, base_path, report_path)
    else:
        make_md_report(project_type, week_number, module_code, tutor_name, title, matriculation, date, base_path, report_path)


def make_word_report(project_type, week_number, module_code, tutor_name, title, matriculation, date, base_path, report_path):
    if project_type == "i":
        doc = make_word("", week_number, module_code, tutor_name, title, matriculation, date)
        save_path = Path(base_path) / f"W{week_number}-Report.docx"
        doc.save(save_path)
    else:
        report_path.mkdir(parents=True, exist_ok=True)
        save_path_ind = Path(report_path) / "Individual-Report.docx"
        save_path_gr = Path(report_path) / "Group-Report.docx"
        doc_ind = make_word("Individual Report", week_number, module_code, tutor_name, title, matriculation, date)
        doc_gr = make_word("Group Report", week_number, module_code, tutor_name, title, matriculation, date)
        doc_ind.save(save_path_ind)
        doc_gr.save(save_path_gr)


def make_word(rep_t, week_number, module_code, tutor_name, title, matriculation, date):
        doc = Document()

        heading = doc.add_heading(
            f"{module_code} W{week_number} Assignment {rep_t}", 1)
        head_p = doc.add_paragraph(
            f"Title: {title}\nTutor: {tutor_name}\nStudent: {matriculation}\nDate: {date}\n")

        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        head_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


        doc.add_page_break()

        doc.add_heading("Introduction", 2)
        doc.add_paragraph(
            "Provide a brief overview of the topic, objectives, and scope of the assignment.")

        doc.add_paragraph("Design")
        doc.add_paragraph(
            "Explain your approach, methodology, and design decisions.")


        if rep_t == "" or rep_t == "Group Report":
            doc.add_paragraph("Testing")
            doc.add_paragraph("Implementation Details")
            doc.add_paragraph("Describe the development and coding process.")

            doc.add_heading("Testing Strategies", 2)
            doc.add_paragraph("Explain how you tested your program.")

            doc.add_paragraph("Test Results")

            table = doc.add_table(rows=4, cols=4)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Test File"
            hdr_cells[1].text = "Description"
            hdr_cells[2].text = "Expected Result"
            hdr_cells[3].text = "Actual Result"

            for i in range(1, 4):
                row_cells = table.rows[i].cells
                row_cells[0].text = f"Test_file_{i}.in"
                row_cells[1].text = ""
                row_cells[2].text = ""
                row_cells[3].text = ""

        doc.add_heading("Conclusion", 2)
        doc.add_paragraph("Summarise key points, pretend you genuinely enjoyed the process of creating this, and talk about potential improvements you'll never implement.")

        doc.add_heading("References (if needed)", 2)
        doc.add_paragraph("Include citations in academic format.")

        return doc


def make_md_report(project_type, week_number, module_code, tutor_name, title, matriculation, date, base_path, report_path):
    # create folder structure
    if project_type == "i":  # individual report setup
        report_file = base_path / f"W{week_number}-Report.md"

        report_type = ""
        report_contents = MD_REPORT_TEMPLATE.substitute(module_code=module_code, week_number=week_number, report_type=report_type, tutor_name=tutor_name, title=title, matriculation=matriculation, date=date)
        # write report files
        report_file.write_text(report_contents)

    else:  # group project setup
        report_path.mkdir(parents=True, exist_ok=True)
        report_type_g = "Group Report"

        individual_report_file = report_path / "Individual-Report.md"
        report_type_i = "Individual Report"

        report_g_contents = MD_REPORT_TEMPLATE.substitute(module_code=module_code, week_number=week_number, report_type=report_type_g, tutor_name=tutor_name, title=title, matriculation=matriculation, date=date)
        group_report_file.write_text(report_g_contents)

        report_i_contents = MD_REPORT_TEMPLATE.substitute(module_code=module_code, week_number=week_number, report_type=report_type_i, tutor_name=tutor_name, title=title, matriculation=matriculation, date=date)
        individual_report_file.write_text(report_i_contents)

def main():
    create_project()

if __name__ == "__main__":
    main()
